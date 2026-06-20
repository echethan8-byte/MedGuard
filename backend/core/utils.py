"""
Core utilities: text extraction, chunking, embedding, ChromaDB management.
"""
import re
import logging
import hashlib
import os
from typing import List, Dict, Any, Optional
from pathlib import Path

logger = logging.getLogger('medguard')


# ─── Text Extraction ──────────────────────────────────────────────────────────

def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from a PDF using PyMuPDF."""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(file_path)
        text_parts = []
        for page_num, page in enumerate(doc):
            text = page.get_text("text")
            if text.strip():
                text_parts.append(f"\n--- Page {page_num + 1} ---\n{text}")
        doc.close()
        return "\n".join(text_parts)
    except Exception as e:
        logger.error(f"PDF extraction failed for {file_path}: {e}")
        raise


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from a DOCX file."""
    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
        return "\n".join(paragraphs)
    except Exception as e:
        logger.error(f"DOCX extraction failed for {file_path}: {e}")
        raise


def extract_text_from_txt(file_path: str) -> str:
    """Extract text from a plain text file."""
    with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
        return f.read()


def extract_text(file_path: str) -> str:
    """Route to correct extractor based on file extension."""
    ext = Path(file_path).suffix.lower()
    if ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif ext == '.docx':
        return extract_text_from_docx(file_path)
    elif ext in ('.txt', '.text'):
        return extract_text_from_txt(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


# ─── Text Cleaning ─────────────────────────────────────────────────────────────

def clean_text(text: str) -> str:
    """Remove boilerplate noise from extracted text."""
    # Collapse excessive whitespace
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r'[ \t]{2,}', ' ', text)
    # Remove common header/footer patterns
    text = re.sub(r'Page \d+ of \d+', '', text, flags=re.IGNORECASE)
    text = re.sub(r'CONFIDENTIAL|DRAFT|FOR INTERNAL USE', '', text, flags=re.IGNORECASE)
    return text.strip()


# ─── Chunking ─────────────────────────────────────────────────────────────────

def chunk_text(
    text: str,
    chunk_size: int = 800,
    chunk_overlap: int = 100,
    min_chunk_size: int = 50,
) -> List[Dict[str, Any]]:
    """
    Split text into overlapping chunks by sentence boundaries.
    Returns list of {text, chunk_index, char_start, char_end}.
    """
    # Split on sentence boundaries
    sentence_endings = re.compile(r'(?<=[.!?])\s+')
    sentences = sentence_endings.split(text)

    chunks = []
    current_chunk = []
    current_size = 0
    chunk_index = 0

    for sentence in sentences:
        sentence = sentence.strip()
        if not sentence:
            continue

        sentence_len = len(sentence)

        # If adding this sentence would exceed chunk_size, finalize current chunk
        if current_size + sentence_len > chunk_size and current_chunk:
            chunk_text_str = ' '.join(current_chunk)
            if len(chunk_text_str) >= min_chunk_size:
                chunks.append({
                    'text': chunk_text_str,
                    'chunk_index': chunk_index,
                })
                chunk_index += 1

            # Overlap: keep last few sentences
            overlap_sentences = []
            overlap_size = 0
            for s in reversed(current_chunk):
                if overlap_size + len(s) < chunk_overlap:
                    overlap_sentences.insert(0, s)
                    overlap_size += len(s)
                else:
                    break
            current_chunk = overlap_sentences
            current_size = overlap_size

        current_chunk.append(sentence)
        current_size += sentence_len

    # Last chunk
    if current_chunk:
        chunk_text_str = ' '.join(current_chunk)
        if len(chunk_text_str) >= min_chunk_size:
            chunks.append({'text': chunk_text_str, 'chunk_index': chunk_index})

    logger.debug(f"Chunked text into {len(chunks)} chunks (size={chunk_size}, overlap={chunk_overlap})")
    return chunks


# ─── Embedding ────────────────────────────────────────────────────────────────

_embedding_model = None

def get_embedding_model(model_name: str = 'all-MiniLM-L6-v2'):
    """Lazy-load sentence transformer model (singleton)."""
    global _embedding_model
    if _embedding_model is None:
        os.environ.setdefault('HF_HUB_OFFLINE', '1')
        os.environ.setdefault('TRANSFORMERS_OFFLINE', '1')
        from sentence_transformers import SentenceTransformer
        logger.info(f"Loading embedding model: {model_name}")
        try:
            _embedding_model = SentenceTransformer(model_name, local_files_only=True)
        except TypeError:
            _embedding_model = SentenceTransformer(model_name)
    return _embedding_model


def embed_texts(texts: List[str], model_name: str = 'all-MiniLM-L6-v2', batch_size: int = 32) -> List[List[float]]:
    """Embed a list of strings. Returns list of float vectors."""
    model = get_embedding_model(model_name)
    embeddings = model.encode(texts, batch_size=batch_size, show_progress_bar=False, convert_to_numpy=True)
    return embeddings.tolist()


# ─── ChromaDB ─────────────────────────────────────────────────────────────────

_chroma_client = None

def get_chroma_client():
    """Lazy-load ChromaDB persistent client."""
    global _chroma_client
    if _chroma_client is None:
        import chromadb
        from django.conf import settings
        _chroma_client = chromadb.PersistentClient(path=settings.CHROMA_PERSIST_DIR)
        logger.info(f"ChromaDB initialized at {settings.CHROMA_PERSIST_DIR}")
    return _chroma_client


def get_or_create_collection(collection_name: str):
    """Get or create a ChromaDB collection."""
    client = get_chroma_client()
    return client.get_or_create_collection(
        name=collection_name,
        metadata={"hnsw:space": "cosine"},
    )


def add_chunks_to_chroma(
    collection_name: str,
    chunks: List[Dict[str, Any]],
    doc_id: str,
    doc_metadata: Dict[str, Any],
    embedding_model: str = 'all-MiniLM-L6-v2',
) -> int:
    """Embed and add chunks to ChromaDB. Returns number of chunks added."""
    if not chunks:
        return 0

    collection = get_or_create_collection(collection_name)
    texts = [c['text'] for c in chunks]
    embeddings = embed_texts(texts, model_name=embedding_model)

    ids = []
    metadatas = []
    for i, chunk in enumerate(chunks):
        chunk_id = hashlib.md5(f"{doc_id}_{i}_{chunk['text'][:50]}".encode()).hexdigest()
        ids.append(chunk_id)
        metadatas.append({
            **doc_metadata,
            'chunk_index': chunk['chunk_index'],
            'doc_id': str(doc_id),
        })

    collection.upsert(ids=ids, embeddings=embeddings, documents=texts, metadatas=metadatas)
    logger.info(f"Added {len(chunks)} chunks to collection '{collection_name}' for doc_id={doc_id}")
    return len(chunks)


def delete_doc_from_chroma(collection_name: str, doc_id: str):
    """Delete all chunks for a document from ChromaDB."""
    collection = get_or_create_collection(collection_name)
    results = collection.get(where={"doc_id": str(doc_id)})
    if results['ids']:
        collection.delete(ids=results['ids'])
        logger.info(f"Deleted {len(results['ids'])} chunks for doc_id={doc_id} from '{collection_name}'")


def similarity_search(
    collection_name: str,
    query: str,
    k: int = 20,
    where: Optional[Dict] = None,
    embedding_model: str = 'all-MiniLM-L6-v2',
) -> List[Dict[str, Any]]:
    """Query ChromaDB and return top-k results with scores."""
    collection = get_or_create_collection(collection_name)
    query_embedding = embed_texts([query], model_name=embedding_model)[0]

    kwargs = {'query_embeddings': [query_embedding], 'n_results': k, 'include': ['documents', 'metadatas', 'distances']}
    if where:
        kwargs['where'] = where

    results = collection.query(**kwargs)

    output = []
    docs = results['documents'][0]
    metas = results['metadatas'][0]
    dists = results['distances'][0]

    for doc, meta, dist in zip(docs, metas, dists):
        output.append({
            'text': doc,
            'metadata': meta,
            'score': 1 - dist,  # Convert cosine distance → similarity
        })

    return output
